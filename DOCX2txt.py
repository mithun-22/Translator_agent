import io
import os
import logging
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Union, IO, Dict, Any

try:
    from .utils import translation_manager
except Exception:
    try:
        from utils import translation_manager
    except Exception:
        translation_manager = None

logger = logging.getLogger(__name__)

class DOCXProcessor:
    def __init__(self, debug_mode: bool = False):
        self.debug_mode = debug_mode

    def process_docx(self, docx_file: Union[str, IO], source_lang: str, target_lang: str, engine: str) -> bytes:
        """
        Translates a DOCX file while preserving layout and formatting.
        """
        try:
            if isinstance(docx_file, (str, bytes)):
                doc = Document(io.BytesIO(docx_file) if isinstance(docx_file, bytes) else docx_file)
            else:
                docx_file.seek(0)
                doc = Document(docx_file)

            # Translate paragraphs
            for para in doc.paragraphs:
                self._translate_paragraph(para, source_lang, target_lang, engine)

            # Translate tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            self._translate_paragraph(para, source_lang, target_lang, engine)

            # Translate headers and footers
            for section in doc.sections:
                for header in [section.header, section.first_page_header, section.even_page_header]:
                    if header:
                        for para in header.paragraphs:
                            self._translate_paragraph(para, source_lang, target_lang, engine)
                for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                    if footer:
                        for para in footer.paragraphs:
                            self._translate_paragraph(para, source_lang, target_lang, engine)

            output = io.BytesIO()
            doc.save(output)
            return output.getvalue()

        except Exception as e:
            logger.error(f"DOCX translation failed: {e}", exc_info=True)
            raise

    def _translate_paragraph(self, para, source_lang, target_lang, engine):
        if not para.text.strip():
            return

        # Simple approach: translate the whole paragraph text and then try to re-apply runs
        # A more complex approach would translate run by run, but that breaks context.
        # So we translate the whole text and put it back into the first run while clearing others.
        
        original_text = para.text
        try:
            res = translation_manager.translate_text(original_text, source_lang, target_lang, engine) if translation_manager else None
            translated_text = res.translated_text if res and res.success else original_text
        except Exception as e:
            logger.warning(f"Paragraph translation failed: {e}")
            translated_text = original_text

        if not para.runs:
            para.add_run(translated_text)
            return

        # Keep track of formatting from the first significant run
        first_run = para.runs[0]
        font_name = first_run.font.name
        font_size = first_run.font.size
        bold = first_run.bold
        italic = first_run.italic
        underline = first_run.underline
        color = first_run.font.color.rgb if first_run.font.color else None

        # Clear runs and add a single translated run
        # NOTE: This loses mid-paragraph formatting changes, but preserves context for translation.
        # For a true 1:1, we'd need to map translated fragments back to runs, which is very hard.
        p_element = para._p
        for run in para.runs:
            p_element.remove(run._r)
        
        new_run = para.add_run(translated_text)
        new_run.font.name = font_name
        new_run.font.size = font_size
        new_run.bold = bold
        new_run.italic = italic
        new_run.underline = underline
        if color:
            new_run.font.color.rgb = color

def translate_docx(docx_file: Union[str, IO], source_lang: str = "en",
                  target_lang: str = "hi", engine: str = "gemini") -> bytes:
    processor = DOCXProcessor()
    return processor.process_docx(docx_file, source_lang, target_lang, engine)
